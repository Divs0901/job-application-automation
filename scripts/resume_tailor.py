"""
resume_tailor.py
────────────────
Tailors your base resume to a job description using Groq AI,
saves the new .docx, and logs everything to the Excel tracker.
"""

import argparse
import os
import re
import json
import datetime
from pathlib import Path

from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl


def read_resume_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_job_desc(job_input: str) -> str:
    if os.path.isfile(job_input):
        return Path(job_input).read_text(encoding="utf-8")
    return job_input


def tailor_resume_with_groq(
    resume_text: str,
    job_desc: str,
    company: str,
    title: str,
    api_key: str,
) -> dict:
    client = Groq(api_key=api_key)

    prompt = f"""You are an expert resume writer and ATS optimization specialist.

TASK: Tailor the resume below to match the job description.
Return ONLY a valid JSON object — no markdown, no explanation, no code fences.

JOB: {title} at {company}

JOB DESCRIPTION:
{job_desc}

CURRENT RESUME:
{resume_text}

Return this exact JSON structure:
{{
  "summary": "2-3 sentence tailored professional summary",
  "skills": ["skill1", "skill2"],
  "experience": [
    {{
      "company": "Company Name",
      "title": "Job Title",
      "dates": "Month Year - Month Year",
      "bullets": [
        "Achievement bullet tailored to job keywords"
      ]
    }}
  ],
  "education": [
    {{
      "school": "University Name",
      "degree": "Degree, Major",
      "dates": "Year"
    }}
  ],
  "keywords_added": ["list", "of", "keywords", "injected"]
}}

Rules:
- Keep all real experience, never invent companies or roles
- Reorder and rephrase bullet points to match job keywords
- Prioritize achievements with metrics (numbers, %)
- Inject keywords from the job description naturally
- Skills section should list top 12-15 most relevant skills
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=4096,
    )

    raw = response.choices[0].message.content.strip()
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"^```\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def build_resume_docx(tailored: dict, candidate_name: str, output_path: str):
    doc = Document()

    # ── Contact Info (hardcoded) ──────────────────────────────────────────────
    CONTACT_EMAIL   = "divyahanda2004@gmail.com"
    CONTACT_PHONE   = "+91 9205574668"
    CONTACT_LINKEDIN = "www.linkedin.com/in/divya-handa-6b84b5212"

    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(54)
        section.right_margin  = Pt(54)

    def heading(text, size=14, bold=True, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        return p

    def divider():
        p = doc.add_paragraph("─" * 85)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.runs[0]
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(8)

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
        return p

    # Name
    name_para = doc.add_paragraph()
    name_run  = name_para.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(31, 56, 100)
    name_para.alignment = 1

    # Contact info line
    contact_para = doc.add_paragraph()
    contact_para.alignment = 1
    contact_para.paragraph_format.space_before = Pt(2)
    contact_para.paragraph_format.space_after  = Pt(6)

    def add_contact(text, is_link=False, add_separator=True):
        run = contact_para.add_run(text)
        run.font.size = Pt(10)
        if is_link:
            run.font.color.rgb = RGBColor(46, 117, 182)
            run.underline = True
        else:
            run.font.color.rgb = RGBColor(80, 80, 80)
        if add_separator:
            sep = contact_para.add_run("  |  ")
            sep.font.size = Pt(10)
            sep.font.color.rgb = RGBColor(180, 180, 180)

    add_contact(CONTACT_EMAIL)
    add_contact(CONTACT_PHONE)
    add_contact(CONTACT_LINKEDIN, is_link=True, add_separator=False)

    # Summary
    heading("PROFESSIONAL SUMMARY", size=11, color=(31, 56, 100))
    divider()
    body(tailored.get("summary", ""))

    # Skills
    heading("CORE COMPETENCIES", size=11, color=(31, 56, 100))
    divider()
    skills = tailored.get("skills", [])
    rows = [skills[i:i+3] for i in range(0, len(skills), 3)]
    for row in rows:
        body("    •  ".join(row))

    # Experience
    heading("PROFESSIONAL EXPERIENCE", size=11, color=(31, 56, 100))
    divider()
    for exp in tailored.get("experience", []):
        p = doc.add_paragraph()
        title_run = p.add_run(f"{exp.get('title', '')}  |  ")
        title_run.bold = True
        title_run.font.size = Pt(11)
        company_run = p.add_run(exp.get("company", ""))
        company_run.bold = True
        company_run.font.size = Pt(11)
        company_run.font.color.rgb = RGBColor(46, 117, 182)

        dates_p = doc.add_paragraph(exp.get("dates", ""))
        dates_p.paragraph_format.space_before = Pt(0)
        dates_p.paragraph_format.space_after  = Pt(2)
        for run in dates_p.runs:
            run.font.size      = Pt(9)
            run.font.italic    = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            run = bp.add_run(bullet)
            run.font.size = Pt(10)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after  = Pt(1)
            bp.paragraph_format.left_indent  = Pt(12)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Education
    heading("EDUCATION", size=11, color=(31, 56, 100))
    divider()
    for edu in tailored.get("education", []):
        p = doc.add_paragraph()
        deg_run = p.add_run(f"{edu.get('degree', '')}  |  ")
        deg_run.bold = True
        deg_run.font.size = Pt(10)
        sch_run = p.add_run(edu.get("school", ""))
        sch_run.font.size = Pt(10)
        sch_run.font.color.rgb = RGBColor(46, 117, 182)
        dates_p = doc.add_paragraph(edu.get("dates", ""))
        for run in dates_p.runs:
            run.font.size   = Pt(9)
            run.font.italic = True

    doc.save(output_path)
    print(f"✅ Tailored resume saved: {output_path}")


def log_to_tracker(tracker_path: str, entry: dict):
    wb = openpyxl.load_workbook(tracker_path)
    ws = wb["Applications"]

    next_row = 2
    for row in ws.iter_rows(min_row=2, max_row=200, min_col=3, max_col=3):
        cell = row[0]
        if not cell.value:
            next_row = cell.row
            break

    ALT   = "EEF2F7"
    WHITE = "FFFFFF"
    fill_color = ALT if next_row % 2 == 0 else WHITE
    fill = openpyxl.styles.PatternFill("solid", start_color=fill_color)
    font = openpyxl.styles.Font(name="Arial", size=10)

    row_data = [
        "",
        entry.get("date", datetime.date.today().strftime("%Y-%m-%d")),
        entry.get("company", ""),
        entry.get("title", ""),
        entry.get("platform", ""),
        entry.get("location", ""),
        entry.get("salary", ""),
        entry.get("status", "Applied"),
        entry.get("resume_version", ""),
        entry.get("resume_path", ""),
        entry.get("job_url", ""),
        entry.get("job_desc_snippet", "")[:500],
        "",
        "",
        "",
        "",
        "",
        "",
        entry.get("notes", ""),
    ]

    for col_idx, value in enumerate(row_data, start=1):
        cell = ws.cell(row=next_row, column=col_idx, value=value)
        cell.fill      = fill
        cell.font      = font
        cell.alignment = openpyxl.styles.Alignment(vertical="center")

    if entry.get("resume_path"):
        resume_cell           = ws.cell(row=next_row, column=10)
        resume_cell.hyperlink = entry["resume_path"]
        resume_cell.value     = Path(entry["resume_path"]).name
        resume_cell.font      = openpyxl.styles.Font(
            name="Arial", size=10, color="2E75B6", underline="single"
        )

    if entry.get("job_url"):
        url_cell           = ws.cell(row=next_row, column=11)
        url_cell.hyperlink = entry["job_url"]
        url_cell.font      = openpyxl.styles.Font(
            name="Arial", size=10, color="2E75B6", underline="single"
        )

    wb.save(tracker_path)
    print(f"✅ Logged to tracker: row {next_row}")


def main():
    parser = argparse.ArgumentParser(description="Tailor resume and log application")
    parser.add_argument("--resume",   required=True)
    parser.add_argument("--job-desc", required=True)
    parser.add_argument("--company",  required=True)
    parser.add_argument("--title",    required=True)
    parser.add_argument("--platform", default="")
    parser.add_argument("--url",      default="")
    parser.add_argument("--location", default="")
    parser.add_argument("--salary",   default="")
    parser.add_argument("--tracker",  required=True)
    parser.add_argument("--api-key",  default=os.getenv("GROQ_API_KEY"))
    parser.add_argument("--name",     default="Your Name")
    parser.add_argument("--out-dir",  default="./resumes")
    args = parser.parse_args()

    if not args.api_key:
        raise ValueError("Set GROQ_API_KEY env var or pass --api-key")

    resume_text = read_resume_text(args.resume)
    job_desc    = read_job_desc(args.job_desc)

    print(f"🤖 Tailoring resume for {args.title} at {args.company}...")
    tailored = tailor_resume_with_groq(
        resume_text, job_desc, args.company, args.title, args.api_key
    )

    os.makedirs(args.out_dir, exist_ok=True)
    safe_company    = re.sub(r"[^\w]", "_", args.company)
    safe_title      = re.sub(r"[^\w]", "_", args.title)
    date_str        = datetime.date.today().strftime("%Y%m%d")
    resume_filename = f"Resume_{safe_company}_{safe_title}_{date_str}.docx"
    resume_path     = os.path.join(args.out_dir, resume_filename)

    build_resume_docx(tailored, args.name, resume_path)

    log_to_tracker(args.tracker, {
        "date":             datetime.date.today().strftime("%Y-%m-%d"),
        "company":          args.company,
        "title":            args.title,
        "platform":         args.platform,
        "location":         args.location,
        "salary":           args.salary,
        "status":           "Applied",
        "resume_version":   resume_filename,
        "resume_path":      os.path.abspath(resume_path),
        "job_url":          args.url,
        "job_desc_snippet": job_desc,
        "notes":            f"Keywords added: {', '.join(tailored.get('keywords_added', []))}"
    })

    print("\n🎉 Done!")
    print(f"   Resume:  {resume_path}")
    print(f"   Tracker: {args.tracker}")
    print(f"   Keywords: {', '.join(tailored.get('keywords_added', []))}")


if __name__ == "__main__":
    main()
